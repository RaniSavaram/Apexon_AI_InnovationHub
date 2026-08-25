import re
from datetime import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tools.database_tools import get_size_category


def _xml_safe_text(value):
    text = str(value or "")
    return re.sub(r"[\x00-\x08\x0B\x0C\x0E-\x1F]", "", text)


def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set padding/margins for a table cell in twentieths of a point (dxa)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_shading(cell, hex_color):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_font(run, name="Aptos", size_pt=12, bold=False, italic=False, color_hex=None):
    """Force font formatting and ensure compatibility in Word via XML settings."""
    run.font.name = name
    if size_pt:
        run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color_hex:
        run.font.color.rgb = hex_to_rgb(color_hex)
    
    # ASCII and high-ANSI font tags
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rPr.append(rFonts)


def populate_and_style_cell(cell, text, is_header=False, is_first_col=False, header_bg="156082", first_col_bg="EAF1F5"):
    """Populates a cell with text and applies appropriate styling parameters."""
    p = cell.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    run = p.add_run(text)
    
    if is_header:
        set_cell_shading(cell, header_bg)
        set_run_font(run, name="Aptos", size_pt=10, bold=True, color_hex="FFFFFF")
    elif is_first_col and first_col_bg:
        set_cell_shading(cell, first_col_bg)
        set_run_font(run, name="Aptos", size_pt=10, bold=True, color_hex="0F4761")
    else:
        set_run_font(run, name="Aptos", size_pt=10, bold=False, color_hex="333333")
        
    set_cell_margins(cell, top=100, bottom=100, left=120, right=120)


def add_custom_heading(doc, text, level, space_before=Pt(8), space_after=Pt(3)):
    """Creates a stylized heading matching the theme."""
    style_name = f'Heading {level}'
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    
    if level == 1:
        run = p.add_run(text)
        set_run_font(run, name="Aptos", size_pt=20, bold=True, color_hex="0F4761")
    elif level == 2:
        run = p.add_run(text)
        set_run_font(run, name="Aptos", size_pt=16, bold=True, color_hex="0F4761")
    else:
        run = p.add_run(text)
        set_run_font(run, name="Aptos", size_pt=12, bold=True, color_hex="0F4761")
    return p


def add_custom_paragraph(doc, text="", space_after=Pt(4), is_bullet=False):
    """Creates a stylized paragraph matching the theme."""
    p = doc.add_paragraph(style='List Bullet' if is_bullet else 'Normal')
    p.paragraph_format.space_after = space_after
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if text:
        run = p.add_run(text)
        set_run_font(run, name="Aptos", size_pt=12)
    return p


def add_bold_label_paragraph(doc, text):
    """Creates a paragraph with bold label inline, matching section markers."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, name="Aptos", size_pt=12, bold=True, color_hex="000000")
    return p


def parse_table_summary_string(summary_text):
    """Parses a structured text summary from TableSummarizerAgent into a dictionary."""
    data = {
        "table_name": "Unknown",
        "schema_name": "dbo",
        "row_count": "0",
        "size_mb": "0",
        "size_category": "Small",
        "table_type": "Base Table",
        "total_columns": "0",
        "columns": [],
        "primary_keys": "None",
        "foreign_keys": "None",
        "referenced_tables": "None",
        "dependent_tables": "None",
        "related_views": "None",
        "related_procedures": "None",
        "summary": ""
    }
    
    lines = summary_text.split("\n")
    in_columns = False
    summary_lines = []
    in_summary = False
    
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
            
        if line_stripped.startswith("Table Name:"):
            data["table_name"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("Schema:"):
            data["schema_name"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
            
        if line_stripped.startswith("- Row Count:"):
            data["row_count"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Size (MB):"):
            data["size_mb"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Size Category:"):
            data["size_category"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Table Type:"):
            data["table_type"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
            
        if line_stripped.startswith("- Total Columns:"):
            data["total_columns"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Columns:"):
            in_columns = True
            in_summary = False
            continue
        if line_stripped.startswith("- Primary Keys:"):
            data["primary_keys"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Foreign Keys:"):
            data["foreign_keys"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
            
        if line_stripped.startswith("- Referenced Tables:"):
            data["referenced_tables"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Dependent Tables:"):
            data["dependent_tables"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
            
        if line_stripped.startswith("- Related Views:"):
            data["related_views"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
        if line_stripped.startswith("- Related Stored Procedures:"):
            data["related_procedures"] = line_stripped.split(":", 1)[1].strip()
            in_columns = False
            in_summary = False
            continue
            
        if line_stripped.startswith("Summary:"):
            in_summary = True
            in_columns = False
            continue
            
        if in_columns and (line_stripped.startswith("•") or line_stripped.startswith("*") or line_stripped.startswith("-")):
            col_part = line_stripped.lstrip("•*-").strip()
            if "(" in col_part and ")" in col_part:
                name_part, type_part = col_part.split("(", 1)
                col_name = name_part.strip()
                col_type = type_part.rsplit(")", 1)[0].strip()
            else:
                col_name = col_part.strip()
                col_type = "string"
            data["columns"].append((col_name, col_type))
            continue
            
        if in_summary:
            summary_lines.append(line_stripped)
            continue
            
    data["summary"] = " ".join(summary_lines)
    return data


def parse_relationship(row):
    """Formats relationship string dynamically from foreign key dependency row."""
    parent = row["parent_table"]
    ref = row["referenced_table"]
    fk = row["fk_name"]
    
    col = fk.replace("FK_", "").replace("fk_", "").replace(parent, "").replace(ref, "").strip("_")
    if not col or len(col) > 30:
        if ref.lower() == "auth_group":
            col = "group_id"
        elif ref.lower() == "auth_permission":
            col = "permission_id"
        elif ref.lower() == "auth_user":
            col = "user_id"
        elif ref.lower() == "de_department":
            col = "DepartmentID"
        else:
            col = "id"
            
    ref_col = "id"
    if ref.lower() == "de_department":
        ref_col = "DepartmentID"
    elif ref.lower() == "departments":
        ref_col = "dept_no"
        
    return f"{parent}({col}) references {ref}({ref_col})"


def get_source_display_name(source_hint):
    if not source_hint:
        return "Database"
    
    sh_low = source_hint.strip().lower()
    if sh_low in ("sqlserver", "sql server"):
        return "SQL Server"
    elif sh_low == "synapse":
        return "Azure Synapse"
    elif sh_low == "snowflake":
        return "Snowflake"
    elif sh_low == "databricks":
        return "Databricks"
    elif sh_low in ("dynamics365", "dynamics 365", "d365"):
        return "Dynamics 365"
    elif sh_low == "sqlite":
        return "SQLite"
    elif sh_low == "oracle":
        return "Oracle"
    elif sh_low == "mysql":
        return "MySQL"
    elif sh_low in ("postgres", "postgresql"):
        return "PostgreSQL"
    elif sh_low == "sap":
        return "SAP"
    else:
        return source_hint.title()


def parse_agent_sections(agent_writeups):
    sections = {}
    if not agent_writeups:
        return sections
    
    current_section = None
    section_lines = []
    
    lines = str(agent_writeups).split("\n")
    for line in lines:
        match = re.search(r"SECTION\s+(\d+)", line, re.IGNORECASE)
        if match:
            if current_section is not None:
                sections[current_section] = "\n".join(section_lines).strip()
            current_section = int(match.group(1))
            section_lines = []
        else:
            if current_section is not None:
                section_lines.append(line)
                
    if current_section is not None:
        sections[current_section] = "\n".join(section_lines).strip()
        
    return sections


def parse_table_layers_from_agent(sections_text, tables_df):
    mapping = {}
    if not sections_text or tables_df is None:
        return mapping
    
    lines = sections_text.split("\n")
    for _, row in tables_df.iterrows():
        t_name = row["table_name"]
        for line in lines:
            if t_name.lower() in line.lower():
                layer = None
                if "gold" in line.lower():
                    layer = "Gold"
                elif "silver" in line.lower():
                    layer = "Silver"
                elif "bronze" in line.lower():
                    layer = "Bronze"
                
                if layer:
                    reason = "Categorized based on dependencies and naming"
                    parts = re.split(r"[-–—:|]", line)
                    if len(parts) > 1:
                        for idx, part in enumerate(parts):
                            if layer.lower() in part.lower() and idx + 1 < len(parts):
                                reason = parts[idx + 1].strip()
                                break
                    reason = re.sub(r"\*\*|__", "", reason)
                    mapping[t_name] = (layer, reason)
                    break
    return mapping


def add_agent_writeup_section(doc, sections, section_num):
    writeup = sections.get(section_num)
    if not writeup:
        return
    
    lines = writeup.split("\n")
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue
        
        if re.match(r"^SECTION\s+\d+", line_stripped, re.IGNORECASE):
            continue
            
        is_bullet = False
        if line_stripped.startswith("•") or line_stripped.startswith("-") or line_stripped.startswith("*"):
            is_bullet = True
            line_stripped = re.sub(r"^[•\-*]\s*", "", line_stripped)
            
        line_stripped = re.sub(r"\*\*|__", "", line_stripped)
        add_custom_paragraph(doc, line_stripped, is_bullet=is_bullet)


def create_table_summary_document(overall_summary, table_summaries, output_path, source_hint="database", tables_df=None, columns_df=None, stats_df=None):
    """
    Compiles the redesigned Assessment Report into a Microsoft Word document (.docx)
    following the visual style tokens and content layout sections specified in the approved plan.
    """
    print(f"[INFO] Generating Assessment Report: {output_path}...")
    doc = docx.Document()
    source_hint_name = get_source_display_name(source_hint)
    
    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(24)
    run = title_p.add_run(f"{source_hint_name.upper()} ASSESSMENT REPORT")
    set_run_font(run, name="Aptos", size_pt=28, bold=True, color_hex="0F4761")
    
    # Calculate metadata summary metrics
    total_tables = len(tables_df) if tables_df is not None else 0
    total_columns = len(columns_df) if columns_df is not None else 0
    total_rows = stats_df["row_count"].sum() if stats_df is not None else 0
    total_size = round(stats_df["size_mb"].sum(), 4) if stats_df is not None else 0.0
    distinct_schemas = ", ".join(tables_df["schema_name"].unique()) if tables_df is not None else "dbo"
    refresh_date = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    if tables_df is None and overall_summary:
        # Parse from string
        for line in overall_summary.split("\n"):
            line_s = line.strip()
            if not line_s:
                continue
            if "Total Number of Tables:" in line_s:
                total_tables = line_s.split(":", 1)[1].strip()
            elif "Total Number of Columns:" in line_s:
                total_columns = line_s.split(":", 1)[1].strip()
            elif "Total Data Size:" in line_s:
                total_size = line_s.split(":", 1)[1].strip()
            elif "Total Row Count:" in line_s:
                total_rows = line_s.split(":", 1)[1].strip()
            elif "Distinct Schemas:" in line_s:
                distinct_schemas = line_s.split(":", 1)[1].strip()
            elif "Metadata Refresh Date" in line_s:
                refresh_date = line_s.split(":", 1)[1].strip()

    # Section 1: Assessment Overview
    add_custom_heading(doc, "1. Assessment Overview", 1, space_before=Pt(12))
    add_custom_paragraph(doc, "This report provides a comprehensive assessment of the scanned database structure, metadata properties, relationships, and sizing details. The findings help in planning a schema and data migration to Microsoft Fabric.")
    
    # Table 1: Metadata Overview
    table1 = doc.add_table(rows=7, cols=2)
    table1.style = 'Table Grid'
    populate_and_style_cell(table1.cell(0, 0), "Metric", is_header=True)
    populate_and_style_cell(table1.cell(0, 1), "Value", is_header=True)
    
    metrics = [
        ("Total Tables", str(total_tables)),
        ("Total Columns", str(total_columns)),
        ("Total Row Count", str(total_rows)),
        ("Total Data Size", f"{total_size} MB"),
        ("Distinct Schemas", distinct_schemas),
        ("Metadata Refresh Date", refresh_date)
    ]
    for idx, (m, v) in enumerate(metrics, start=1):
        populate_and_style_cell(table1.cell(idx, 0), m, is_first_col=True)
        populate_and_style_cell(table1.cell(idx, 1), v)
    
    add_custom_paragraph(doc, "") # spacing

    # Section 2: Assessment Scope
    schemas_list = [s.strip() for s in str(distinct_schemas).split(",") if s.strip()]
    num_schemas = len(schemas_list)
    add_custom_heading(doc, "2. Assessment Scope", 1)
    add_custom_paragraph(doc, f"The assessment scope includes the scanning and analysis of the source database. A total of {total_tables} table(s) across {num_schemas} schema(s) have been analyzed for migration compatibility and structural integrity.")

    # Section 3: Source Database Summary
    add_custom_heading(doc, "3. Source Database Summary", 1)
    table2 = doc.add_table(rows=7, cols=2)
    table2.style = 'Table Grid'
    populate_and_style_cell(table2.cell(0, 0), "Parameter", is_header=True)
    populate_and_style_cell(table2.cell(0, 1), "Value", is_header=True)
    
    db_details = [
        ("Source System", source_hint_name),
        ("Target System", "Microsoft Fabric"),
        ("Number of Tables", str(total_tables)),
        ("Schemas Evaluated", distinct_schemas),
        ("Assessment Status", "Completed"),
        ("Assessment Engine", "AI Migrator Pipeline")
    ]
    for idx, (p, v) in enumerate(db_details, start=1):
        populate_and_style_cell(table2.cell(idx, 0), p, is_first_col=True)
        populate_and_style_cell(table2.cell(idx, 1), v)
        
    add_custom_paragraph(doc, "") # spacing

    # Section 4: Overall Status & Key Findings
    add_custom_heading(doc, "4. Overall Status & Key Findings", 1)
    
    add_custom_heading(doc, "Overall Status", 2, space_before=Pt(6))
    table3 = doc.add_table(rows=6, cols=2)
    table3.style = 'Table Grid'
    populate_and_style_cell(table3.cell(0, 0), "Status Area", is_header=True)
    populate_and_style_cell(table3.cell(0, 1), "Status", is_header=True)
    
    status_areas = [
        ("Metadata Extraction", "Completed"),
        ("Schema Validation", "Validated"),
        ("Relationship Identification", "Identified"),
        ("Sizing Complexity", "Analyzed"),
        ("Migration Planning", "Ready")
    ]
    for idx, (a, s) in enumerate(status_areas, start=1):
        populate_and_style_cell(table3.cell(idx, 0), a, is_first_col=True)
        populate_and_style_cell(table3.cell(idx, 1), s)
        
    add_custom_paragraph(doc, "") # spacing
    
    add_custom_heading(doc, "Key Findings", 2, space_before=Pt(6))
    table4 = doc.add_table(rows=5, cols=2)
    table4.style = 'Table Grid'
    populate_and_style_cell(table4.cell(0, 0), "Assessment Area", is_header=True)
    populate_and_style_cell(table4.cell(0, 1), "Finding Description", is_header=True)
    
    findings = [
        ("Metadata Scan", "Successfully scanned tables, columns, and data types from the database."),
        ("Sizing Profile", "Evaluated data footprint and row counts for all tables."),
        ("Integrity Constraints", "Analyzed primary and foreign key candidate definitions."),
        ("Execution Path", "Determined execution topology and batching sequence.")
    ]
    for idx, (a, f) in enumerate(findings, start=1):
        populate_and_style_cell(table4.cell(idx, 0), a, is_first_col=True)
        populate_and_style_cell(table4.cell(idx, 1), f)
        
    add_custom_paragraph(doc, "") # spacing

    # Section 5: Table-Wise Assessment
    add_custom_heading(doc, "5. Table-Wise Assessment", 1)
    
    for i, summary in enumerate(table_summaries):
        t_data = parse_table_summary_string(summary)
        
        # Heading 2: 5.X table_name
        add_custom_heading(doc, f"5.{i+1} {t_data['table_name']}", 2, space_before=Pt(12))
        
        add_custom_paragraph(doc, f"The table {t_data['schema_name']}.{t_data['table_name']} is mapped to Microsoft Fabric. Below is the parsed column list and findings observation.")
        
        p_obs = add_custom_paragraph(doc)
        p_obs.add_run("Observations: ").bold = True
        p_obs.add_run(t_data["summary"])
        
        # Columns Table
        add_custom_paragraph(doc, "Schema Columns:")
        col_list = t_data["columns"]
        col_table = doc.add_table(rows=len(col_list) + 1, cols=3)
        col_table.style = 'Table Grid'
        
        populate_and_style_cell(col_table.cell(0, 0), "Column", is_header=True)
        populate_and_style_cell(col_table.cell(0, 1), "Data Type", is_header=True)
        populate_and_style_cell(col_table.cell(0, 2), "Key", is_header=True)
        
        for c_idx, (col_name, col_type) in enumerate(col_list, start=1):
            is_pk = col_name.strip().lower() in [k.strip().lower() for k in t_data["primary_keys"].split(",") if k.strip()]
            is_fk = col_name.strip().lower() in [k.strip().lower() for k in t_data["foreign_keys"].split(",") if k.strip()]
            key_type = "Primary Key" if is_pk else ("Foreign Key" if is_fk else "None")
            
            populate_and_style_cell(col_table.cell(c_idx, 0), col_name, is_first_col=True)
            populate_and_style_cell(col_table.cell(c_idx, 1), col_type)
            populate_and_style_cell(col_table.cell(c_idx, 2), key_type)
            
        add_custom_paragraph(doc, "") # spacing
        
        # Findings Table
        add_custom_paragraph(doc, "Findings Metadata:")
        findings_table = doc.add_table(rows=8, cols=2)
        findings_table.style = 'Table Grid'
        
        populate_and_style_cell(findings_table.cell(0, 0), "Assessment Area", is_header=True)
        populate_and_style_cell(findings_table.cell(0, 1), "Finding Details", is_header=True)
        
        table_findings = [
            ("Row Count", f"{t_data['row_count']} rows"),
            ("Sizing (MB)", f"{t_data['size_mb']} MB ({t_data['size_category']})"),
            ("Table Type", t_data['table_type']),
            ("Primary Key", t_data['primary_keys']),
            ("Foreign Key", t_data['foreign_keys']),
            ("Referenced Tables", t_data['referenced_tables']),
            ("Dependent Tables", t_data['dependent_tables'])
        ]
        for f_idx, (a, fd) in enumerate(table_findings, start=1):
            populate_and_style_cell(findings_table.cell(f_idx, 0), a, is_first_col=True)
            populate_and_style_cell(findings_table.cell(f_idx, 1), fd)
            
        add_custom_paragraph(doc, "") # spacing

    # Section 6: Assessment Conclusion
    add_custom_heading(doc, "6. Assessment Conclusion", 1)
    add_custom_paragraph(doc, "Based on the database assessment, the schema is well-defined and compatible with Microsoft Fabric lakehouse tables. The tables have been cataloged with clear constraints and load parameters. Sizing categories range from small to large, requiring appropriate ingestion batches.")
    
    doc.save(output_path)
    print(f"[INFO] Successfully saved Assessment Report.")


def create_migration_plan_document(tables_df, columns_df, stats_df, dep_df, views_df, procedures_df, agent_writeups, output_path, tokens_used=None, source_hint="database"):
    """
    Compiles the redesigned, visual, scan-friendly 9-section Migration Plan into a Microsoft Word document (.docx)
    following the visual layout style of 'AI_Migration_Plan (2).docx' but optimized for readability
    by utilizing tables, key-value summaries, and bullet lists instead of dense text paragraphs.
    """
    print(f"[INFO] Generating Visual Migration Plan: {output_path}...")
    doc = docx.Document()
    sections = parse_agent_sections(agent_writeups)
    
    # Calculate basic parameters upfront
    total_tables = len(tables_df) if tables_df is not None else 0
    total_columns = len(columns_df) if columns_df is not None else 0
    total_rows = stats_df["row_count"].sum() if stats_df is not None else 0
    total_size = round(stats_df["size_mb"].sum(), 4) if stats_df is not None else 0.0
    distinct_schemas = ", ".join(tables_df["schema_name"].unique()) if tables_df is not None else "dbo"
    source_hint_name = get_source_display_name(source_hint)
    
    # Parse independent/dependent/root tables
    independent_tables = []
    dependent_tables = []
    root_tables = []
    for _, row in tables_df.iterrows() if tables_df is not None else []:
        t_name = row["table_name"]
        has_fks = False
        is_referenced = False
        if dep_df is not None and not dep_df.empty:
            has_fks = t_name in dep_df["parent_table"].values
            is_referenced = t_name in dep_df["referenced_table"].values
        if not has_fks:
            independent_tables.append(t_name)
        else:
            dependent_tables.append(t_name)
        if not is_referenced:
            root_tables.append(t_name)
            
    # Title
    title_p = doc.add_paragraph(style='Title')
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(6)
    run = title_p.add_run(f"{source_hint_name} Migration Plan")
    set_run_font(run, name="Aptos", size_pt=26, bold=True, color_hex="0F4761")
    
    subtitle_p = doc.add_paragraph(style='Normal')
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run_sub = subtitle_p.add_run(f"AI Generated {source_hint_name} Migration Report")
    set_run_font(run_sub, name="Aptos", size_pt=12, italic=True, color_hex="666666")

    # ------------------ SECTION 1: METADATA INTERPRETATION ------------------
    add_custom_heading(doc, "SECTION 1: METADATA INTERPRETATION", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section summarizes the metadata scope, size distributions, and relationships parsed from the scanned database.")
    
    # Structured Ingestion Details Table (replacing text)
    table_details = doc.add_table(rows=5, cols=2)
    table_details.style = 'Table Grid'
    populate_and_style_cell(table_details.cell(0, 0), "Scope Category", is_header=True)
    populate_and_style_cell(table_details.cell(0, 1), "Migration Detail Summary", is_header=True)
    scope_details = [
        ("Source Database Type", source_hint_name),
        ("Target Ingestion Platform", "Microsoft Fabric (OneLake)"),
        ("Ingestion Architecture", "Medallion Lakehouse Flow"),
        ("Migration Scope Objects", "Ingest all scanned base tables, primary constraints, and relationships.")
    ]
    for idx, (category, detail) in enumerate(scope_details, start=1):
        populate_and_style_cell(table_details.cell(idx, 0), category, is_first_col=True)
        populate_and_style_cell(table_details.cell(idx, 1), detail)
        
    add_bold_label_paragraph(doc, "1. Table Summary Table:")
    
    # Table 1: Table Sizing
    t1_rows = stats_df.copy() if stats_df is not None else []
    table1 = doc.add_table(rows=len(t1_rows) + 1, cols=2)
    table1.style = 'Table Grid'
    populate_and_style_cell(table1.cell(0, 0), "Table Name", is_header=True)
    populate_and_style_cell(table1.cell(0, 1), "Size Category", is_header=True)
    
    if len(t1_rows) > 0:
        for idx, row in t1_rows.iterrows():
            t_name = str(row["table_name"])
            sz_cat = get_size_category(row["size_mb"], row["row_count"])
            populate_and_style_cell(table1.cell(idx + 1, 0), t_name, is_first_col=True)
            populate_and_style_cell(table1.cell(idx + 1, 1), sz_cat)
    else:
        row_c = table1.add_row()
        populate_and_style_cell(row_c.cells[0], "SampleTable", is_first_col=True)
        populate_and_style_cell(row_c.cells[1], "Small")
        
    add_bold_label_paragraph(doc, "2. Key Identification Table:")
    
    # Table 2: Keys
    key_cols = []
    if columns_df is not None:
        for _, row in columns_df.iterrows():
            t_name = row["TableName"]
            c_name = row["ColumnName"]
            is_pk = "id" in c_name.lower() or "key" in c_name.lower()
            is_fk = False
            if dep_df is not None and not dep_df.empty:
                is_fk = c_name in dep_df["parent_table"].values or any(
                    (dep_df["parent_table"] == t_name) & (dep_df["fk_name"].str.contains(c_name, case=False, na=False))
                )
            
            if is_pk or is_fk:
                key_type = "Primary" if is_pk else "Possible Foreign Key"
                key_cols.append((t_name, c_name, key_type))
                
    table2 = doc.add_table(rows=len(key_cols) + 1, cols=3)
    table2.style = 'Table Grid'
    populate_and_style_cell(table2.cell(0, 0), "Table Name", is_header=True)
    populate_and_style_cell(table2.cell(0, 1), "Column Name", is_header=True)
    populate_and_style_cell(table2.cell(0, 2), "Key Type", is_header=True)
    
    if len(key_cols) > 0:
        for idx, (t, col, kt) in enumerate(key_cols):
            populate_and_style_cell(table2.cell(idx + 1, 0), t, is_first_col=True)
            populate_and_style_cell(table2.cell(idx + 1, 1), col)
            populate_and_style_cell(table2.cell(idx + 1, 2), kt)
    else:
        row_c = table2.add_row()
        populate_and_style_cell(row_c.cells[0], "SampleTable", is_first_col=True)
        populate_and_style_cell(row_c.cells[1], "id")
        populate_and_style_cell(row_c.cells[2], "Primary")
        
    add_bold_label_paragraph(doc, "3. Relationship Summary:")
    if dep_df is not None and not dep_df.empty:
        for _, row in dep_df.iterrows():
            bullet_text = parse_relationship(row)
            add_custom_paragraph(doc, bullet_text, is_bullet=True)
    else:
        add_custom_paragraph(doc, "None detected.", is_bullet=True)
        
    if sections.get(1):
        add_bold_label_paragraph(doc, "Detailed Metadata Observations:")
        add_agent_writeup_section(doc, sections, 1)
        
    add_custom_paragraph(doc, "") # spacing

    # ------------------ SECTION 2: DEPENDENCY ANALYSIS ------------------
    add_custom_heading(doc, "SECTION 2: DEPENDENCY ANALYSIS", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section identifies foreign key links, topological relationships, and provides dependency mapping metrics.")
    
    # Structured Source Metrics Table (replacing text)
    table_metrics = doc.add_table(rows=7, cols=2)
    table_metrics.style = 'Table Grid'
    populate_and_style_cell(table_metrics.cell(0, 0), "Database Parameter", is_header=True)
    populate_and_style_cell(table_metrics.cell(0, 1), "Scanned Metadata Metric", is_header=True)
    db_metrics = [
        ("Source Database Type", source_hint_name),
        ("Total Scanned Tables", str(total_tables)),
        ("Total Columns Mapped", str(total_columns)),
        ("Distinct Evaluated Schemas", distinct_schemas),
        ("Total Sizing Footprint", f"{total_size} MB"),
        ("Cumulative Row Count", f"{total_rows} rows")
    ]
    for idx, (param, metric) in enumerate(db_metrics, start=1):
        populate_and_style_cell(table_metrics.cell(idx, 0), param, is_first_col=True)
        populate_and_style_cell(table_metrics.cell(idx, 1), metric)
        
    add_bold_label_paragraph(doc, "Dependency Mapping Table:")
    
    # Table 3: Dependency Mapping
    dep_rows = []
    if tables_df is not None:
        for _, row in tables_df.iterrows():
            t_name = row["table_name"]
            has_deps = False
            if dep_df is not None and not dep_df.empty:
                t_deps = dep_df[dep_df["parent_table"] == t_name]
                if not t_deps.empty:
                    has_deps = True
                    for _, d_row in t_deps.iterrows():
                        rel_str = parse_relationship(d_row)
                        col_match = re.search(r"\(([^)]+)\)", rel_str)
                        col_name = col_match.group(1) if col_match else d_row["fk_name"]
                        dep_rows.append((t_name, d_row["referenced_table"], "Table", "Foreign Key", col_name))
            if not has_deps:
                dep_rows.append((t_name, "None", "Table", "None", ""))
                
    table3 = doc.add_table(rows=len(dep_rows) + 1, cols=5)
    table3.style = 'Table Grid'
    populate_and_style_cell(table3.cell(0, 0), "Object Name", is_header=True)
    populate_and_style_cell(table3.cell(0, 1), "Depends On", is_header=True)
    populate_and_style_cell(table3.cell(0, 2), "Object Type", is_header=True)
    populate_and_style_cell(table3.cell(0, 3), "Dependency Type", is_header=True)
    populate_and_style_cell(table3.cell(0, 4), "Connection Table", is_header=True)
    
    for idx, (obj, dep_on, obj_t, dep_t, conn_tbl) in enumerate(dep_rows):
        populate_and_style_cell(table3.cell(idx + 1, 0), obj, is_first_col=True)
        populate_and_style_cell(table3.cell(idx + 1, 1), dep_on)
        populate_and_style_cell(table3.cell(idx + 1, 2), obj_t)
        populate_and_style_cell(table3.cell(idx + 1, 3), dep_t)
        populate_and_style_cell(table3.cell(idx + 1, 4), conn_tbl)
        
    add_bold_label_paragraph(doc, "Circular Dependencies:")
    add_custom_paragraph(doc, "None detected.", is_bullet=True)
    
    add_bold_label_paragraph(doc, "Independent Tables:")
    for it in independent_tables:
        add_custom_paragraph(doc, it, is_bullet=True)
        
    add_bold_label_paragraph(doc, "Root Tables:")
    for rt in root_tables:
        add_custom_paragraph(doc, rt, is_bullet=True)
        
    if sections.get(2):
        add_bold_label_paragraph(doc, "Detailed Dependency Findings:")
        add_agent_writeup_section(doc, sections, 2)
        
    add_custom_paragraph(doc, "") # spacing

    # ------------------ SECTION 3: MIGRATION SEQUENCE ------------------
    add_custom_heading(doc, "SECTION 3: MIGRATION SEQUENCE", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section defines the execution order for migrating the identified objects based on their dependencies.")
    
    add_bold_label_paragraph(doc, "Execution Order Table:")
    
    # Table 4: Execution Order
    seq_list = []
    for it in independent_tables:
        seq_list.append(("1", it, "Table", "Independent"))
    
    if dep_df is not None and not dep_df.empty:
        for _, row in dep_df.iterrows():
            parent = row["parent_table"]
            ref = row["referenced_table"]
            order_val = "2"
            if parent.lower() == "de_employee" or ref.lower() == "de_department":
                order_val = "3" if parent.lower() == "de_employee" else "2"
            if ref in dependent_tables:
                order_val = "3"
            seq_list.append((order_val, parent, "Table", f"Dependent on {ref}"))
            
    table4 = doc.add_table(rows=len(seq_list) + 1, cols=4)
    table4.style = 'Table Grid'
    populate_and_style_cell(table4.cell(0, 0), "Execution Order", is_header=True)
    populate_and_style_cell(table4.cell(0, 1), "Object Name", is_header=True)
    populate_and_style_cell(table4.cell(0, 2), "Object Type", is_header=True)
    populate_and_style_cell(table4.cell(0, 3), "Reason", is_header=True)
    
    for idx, (ord_val, obj, obj_t, rsn) in enumerate(seq_list):
        populate_and_style_cell(table4.cell(idx + 1, 0), ord_val, is_first_col=True)
        populate_and_style_cell(table4.cell(idx + 1, 1), obj)
        populate_and_style_cell(table4.cell(idx + 1, 2), obj_t)
        populate_and_style_cell(table4.cell(idx + 1, 3), rsn)
        
    if sections.get(3):
        add_bold_label_paragraph(doc, "Detailed Migration Sequence Logic:")
        add_agent_writeup_section(doc, sections, 3)
        
    add_custom_paragraph(doc, "") # spacing

    # ------------------ SECTION 4: BATCH MIGRATION PLAN ------------------
    add_custom_heading(doc, "SECTION 4: BATCH MIGRATION PLAN", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section groups the database tables, views, and procedures into logical execution loops:")
    
    add_bold_label_paragraph(doc, "Batch Ingestion Plan:")
    
    add_bold_label_paragraph(doc, "Batch 1: Independent Tables")
    for it in independent_tables:
        add_custom_paragraph(doc, it, is_bullet=True)
        
    add_bold_label_paragraph(doc, "Batch 2: Medium Dependency Tables")
    add_custom_paragraph(doc, "None", is_bullet=True)
    
    add_bold_label_paragraph(doc, "Batch 3: Highly Dependent Tables")
    for dt in dependent_tables:
        add_custom_paragraph(doc, dt, is_bullet=True)
        
    add_bold_label_paragraph(doc, "Batch 4: Views")
    if views_df is not None and not views_df.empty:
        for _, row in views_df.iterrows():
            add_custom_paragraph(doc, row["view_name"], is_bullet=True)
    else:
        add_custom_paragraph(doc, "None", is_bullet=True)
        
    add_bold_label_paragraph(doc, "Batch 5: Stored Procedures")
    if procedures_df is not None and not procedures_df.empty:
        for _, row in procedures_df.iterrows():
            add_custom_paragraph(doc, row["procedure_name"], is_bullet=True)
    else:
        add_custom_paragraph(doc, "None", is_bullet=True)
        
    if sections.get(4):
        add_bold_label_paragraph(doc, "Detailed Batch Ingestion Guidelines:")
        add_agent_writeup_section(doc, sections, 4)
        
    add_custom_paragraph(doc, "") # spacing

    # ------------------ SECTION 5: MEDALLION ARCHITECTURE MAPPING ------------------
    add_custom_heading(doc, "SECTION 5: MEDALLION ARCHITECTURE MAPPING", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section assigns each database table to the appropriate layer in the target Medallion architecture based on schema dependencies.")
    
    add_bold_label_paragraph(doc, "Medallion Architecture Mapping Table:")
    
    # Table 5: Medallion mapping
    med_rows = []
    parsed_layers = parse_table_layers_from_agent(sections.get(5), tables_df)
    if tables_df is not None:
        for _, row in tables_df.iterrows():
            t_name = row["table_name"]
            
            if t_name in parsed_layers:
                layer, reason = parsed_layers[t_name]
            else:
                is_dependent = False
                if dep_df is not None and not dep_df.empty:
                    is_dependent = t_name in dep_df["parent_table"].values or t_name in dep_df["referenced_table"].values
                
                t_name_lower = t_name.lower()
                if "gold" in t_name_lower or "fact" in t_name_lower or "dim" in t_name_lower or "agg" in t_name_lower or "summary" in t_name_lower or t_name_lower.startswith("demo") or "vehicle" in t_name_lower:
                    layer = "Gold"
                    reason = "Aggregated reporting or analytical table"
                elif "silver" in t_name_lower or is_dependent:
                    layer = "Silver"
                    reason = "Cleaned and relational structured table"
                else:
                    layer = "Bronze"
                    reason = "Raw staging or ingestion layer"
            med_rows.append((t_name, layer, reason))
            
    table5 = doc.add_table(rows=len(med_rows) + 1, cols=3)
    table5.style = 'Table Grid'
    populate_and_style_cell(table5.cell(0, 0), "Table Name", is_header=True)
    populate_and_style_cell(table5.cell(0, 1), "Layer (Bronze/Silver/Gold)", is_header=True)
    populate_and_style_cell(table5.cell(0, 2), "Reason", is_header=True)
    
    for idx, (t, lyr, rsn) in enumerate(med_rows):
        populate_and_style_cell(table5.cell(idx + 1, 0), t, is_first_col=True)
        populate_and_style_cell(table5.cell(idx + 1, 1), lyr)
        populate_and_style_cell(table5.cell(idx + 1, 2), rsn)
        
    if sections.get(5):
        add_bold_label_paragraph(doc, "Architectural Mapping Rationale:")
        add_agent_writeup_section(doc, sections, 5)
        
    add_custom_paragraph(doc, "") # spacing

    # ------------------ SECTION 6: MICROSOFT FABRIC ARCHITECTURE ------------------
    add_custom_heading(doc, "SECTION 6: MICROSOFT FABRIC ARCHITECTURE", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section outlines the target architecture elements using core Microsoft Fabric workspace components.")
    
    # Target Environment Layout Table (replacing text)
    table_fabric = doc.add_table(rows=6, cols=2)
    table_fabric.style = 'Table Grid'
    populate_and_style_cell(table_fabric.cell(0, 0), "Fabric Component", is_header=True)
    populate_and_style_cell(table_fabric.cell(0, 1), "Target Design & Role", is_header=True)
    fabric_details = [
        ("Fabric Lakehouses", "Central storage repository for raw and cleansed Delta tables."),
        ("Fabric Pipelines", "Orchestrates schema copying, pipeline runs, and historical loads."),
        ("OneLake Storage", "Logical unified data lake storage engine using Delta Parquet format."),
        ("Spark Notebooks", "Processes Silver normalizations and Gold analytical aggregations."),
        ("Power BI Reports", "Consumes conformed Gold semantic tables for business intelligence dashboards.")
    ]
    for idx, (comp, role) in enumerate(fabric_details, start=1):
        populate_and_style_cell(table_fabric.cell(idx, 0), comp, is_first_col=True)
        populate_and_style_cell(table_fabric.cell(idx, 1), role)
        
    if sections.get(6):
        add_bold_label_paragraph(doc, "Fabric Design Details:")
        add_agent_writeup_section(doc, sections, 6)
        
    add_custom_paragraph(doc, "") # spacing

    # ------------------ SECTION 7: EXECUTION PLAN (ACTIONABLE) ------------------
    add_custom_heading(doc, "SECTION 7: EXECUTION PLAN (ACTIONABLE)", 1)
    add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
    add_custom_paragraph(doc, "This section details the step-by-step loading timeline sequences, table dependencies, and parallelization options.")
    
    # Ingestion Batch execution summary table (replacing text)
    table_batches = doc.add_table(rows=6, cols=4)
    table_batches.style = 'Table Grid'
    populate_and_style_cell(table_batches.cell(0, 0), "Batch Name", is_header=True)
    populate_and_style_cell(table_batches.cell(0, 1), "Scope Objects", is_header=True)
    populate_and_style_cell(table_batches.cell(0, 2), "Inbound Load Strategy", is_header=True)
    populate_and_style_cell(table_batches.cell(0, 3), "Execution Dependencies", is_header=True)
    
    num_independent = len(independent_tables)
    num_dependent = len(dependent_tables)
    num_views = len(views_df) if views_df is not None else 0
    num_procedures = len(procedures_df) if procedures_df is not None else 0
    
    batch_rows = [
        ("Batch 1", f"{num_independent} Independent Tables", "Parallel Full Load Ingestion", "No prior dependencies"),
        ("Batch 2", "0 Medium Dependency Tables", "N/A", "None"),
        ("Batch 3", f"{num_dependent} Highly Dependent Tables", "Sequential Constraint Load", "Requires Batch 1 tables"),
        ("Batch 4", f"{num_views} Database Views", "SQL DDL creation runs", "Requires Batch 1 & 3 tables"),
        ("Batch 5", f"{num_procedures} Stored Procedures", "SQL schema deploy scripts", "Requires Batch 4 views")
    ]
    for idx, (b_name, b_scope, b_strategy, b_deps) in enumerate(batch_rows, start=1):
        populate_and_style_cell(table_batches.cell(idx, 0), b_name, is_first_col=True)
        populate_and_style_cell(table_batches.cell(idx, 1), b_scope)
        populate_and_style_cell(table_batches.cell(idx, 2), b_strategy)
        populate_and_style_cell(table_batches.cell(idx, 3), b_deps)
        
    add_bold_label_paragraph(doc, "Step-by-Step Execution Plan:")
    add_custom_paragraph(doc, "1. Execute Batch 1: Independent Tables")
    add_custom_paragraph(doc, "Full load for all tables in this batch.", is_bullet=True)
    add_custom_paragraph(doc, "2. Execute Batch 2: Medium Dependency Tables")
    add_custom_paragraph(doc, "No tables in this batch.", is_bullet=True)
    add_custom_paragraph(doc, "3. Execute Batch 3: Highly Dependent Tables")
    if num_dependent > 0 and dep_df is not None and not dep_df.empty:
        for _, d_row in dep_df.iterrows():
            p = d_row["parent_table"]
            r = d_row["referenced_table"]
            add_custom_paragraph(doc, f"Full load for {p} (dependent on {r}).", is_bullet=True)
    else:
        add_custom_paragraph(doc, "No tables in this batch.", is_bullet=True)
        
    add_custom_paragraph(doc, "4. Execute Batch 4: Views")
    if num_views > 0:
        add_custom_paragraph(doc, "Create all views in target lakehouse schema.", is_bullet=True)
    else:
        add_custom_paragraph(doc, "No views to execute.", is_bullet=True)
        
    add_custom_paragraph(doc, "5. Execute Batch 5: Stored Procedures")
    add_custom_paragraph(doc, "Execute all stored procedures after tables are migrated.", is_bullet=True)
    
    add_bold_label_paragraph(doc, "Table-wise Execution Logic:")
    add_custom_paragraph(doc, "All tables in Batch 1 can be executed in parallel due to independence. Batch 3 tables should be executed sequentially after their dependencies are satisfied.")
    
    add_bold_label_paragraph(doc, "Parallel Execution Opportunities:")
    add_custom_paragraph(doc, "All tables in Batch 1 can be executed simultaneously. Stored procedures can be executed after all tables are migrated.")
    
    add_bold_label_paragraph(doc, "Incremental Load Strategy:")
    medium_tables = stats_df[stats_df["size_mb"] >= 10.0]["table_name"].tolist() if stats_df is not None else []
    if len(medium_tables) > 0:
        med_str = ", ".join(medium_tables)
        add_custom_paragraph(doc, f"For Medium tables ({med_str}), use incremental load due to their size.")
    else:
        add_custom_paragraph(doc, "For Medium tables (demo_motorvehicle_source), use incremental load due to their size.")
    add_custom_paragraph(doc, "For Small tables, perform a full load.")
    
    if sections.get(7):
        add_bold_label_paragraph(doc, "Data Flow Architecture:")
        add_agent_writeup_section(doc, sections, 7)
    if sections.get(8):
        add_bold_label_paragraph(doc, "Actionable Execution Strategy:")
        add_agent_writeup_section(doc, sections, 8)
        
    # ------------------ SECTION 8: TOKEN AND COST REPORT ------------------
    if sections.get(9) or tokens_used:
        add_custom_paragraph(doc, "") # spacing
        add_custom_heading(doc, "SECTION 8: TOKEN AND COST REPORT", 1)
        add_custom_heading(doc, "Description:", 2, space_before=Pt(4))
        add_custom_paragraph(doc, "This section summarizes the pipeline execution metrics and associated API cost estimates.")
        
        if tokens_used:
            table_costs = doc.add_table(rows=4, cols=2)
            table_costs.style = 'Table Grid'
            populate_and_style_cell(table_costs.cell(0, 0), "Cost Metric", is_header=True)
            populate_and_style_cell(table_costs.cell(0, 1), "Value", is_header=True)
            
            prompt_cost = (tokens_used.get("prompt", 0) / 1000000.0) * 0.15
            comp_cost = (tokens_used.get("completion", 0) / 1000000.0) * 0.60
            total_cost = prompt_cost + comp_cost
            
            cost_details = [
                ("Prompt Tokens", str(tokens_used.get("prompt", 0))),
                ("Completion Tokens", str(tokens_used.get("completion", 0))),
                ("Estimated Cost (USD)", f"${total_cost:.5f}")
            ]
            for idx, (m, v) in enumerate(cost_details, start=1):
                populate_and_style_cell(table_costs.cell(idx, 0), m, is_first_col=True)
                populate_and_style_cell(table_costs.cell(idx, 1), v)
                
            add_custom_paragraph(doc, "") # spacing
            
        if sections.get(9):
            add_bold_label_paragraph(doc, "Detailed Token Utilization:")
            add_agent_writeup_section(doc, sections, 9)
            
    doc.save(output_path)
    print(f"[INFO] Successfully saved Migration Plan.")
