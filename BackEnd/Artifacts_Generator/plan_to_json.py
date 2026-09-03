"""
Parses the AI-generated Assessment Report (Assesment Report.docx /
<source>_Assessment_Report.docx) - and, optionally, the companion AI
Migration Plan docx (AI_Migration_Plan.docx / <source>_Migration_Plan.docx)
- into migration_plan.json: the normalized JSON contract that
DB2_2_Fabric.py (and any future target-specific generator) consumes to
create Delta tables in Microsoft Fabric.

Output shape
------------
{
  "meta": {"source_system": "databricks", "database_name": "sales_prod"},
  "tables": [
    {
      "schema": "cards",
      "table_name": "dim_card",
      "medallion_layer": "Gold",          # or null if no migration plan given
      "load_strategy": "Full Load",       # or null if no migration plan given
      "columns": [
        {"name": "card_id", "data_type": "STRING"},
        ...
      ]
    },
    ...
  ],
  "views": [
    {"schema": "cards", "view_name": "vw_active_cards", "columns": [...], "definition": "SELECT ..."},
    ...
  ],
  "functions": [
    {"schema": "cards", "function_name": "fn_mask_pan", "return_type": "STRING"},
    ...
  ],
  "volumes": [
    {"schema": "cards", "volume_name": "raw_files", "volume_type": "EXTERNAL", "storage_location": "s3://..."},
    ...
  ]
}

Column/schema parsing supports both report layouts the pipeline has
produced over time:
  - "new format" (current docx_generator.py output): a "5.N <name>"
    Heading 2 per object (table, view, function, or volume), a
    "The <table|view|function|volume> <schema>.<name> is mapped..." sentence
    giving the schema and object type, a "Column | Data Type | Key" docx
    table for tables/views, and a "Property | Value" docx table for
    views/functions/volumes (Definition / Return Type / Volume Type /
    Storage Location).
  - "old format" (a plain-text table_summarizer transcript): "Table Name:"
    / "Schema:" paragraphs followed by a "- Columns:" block of "• name
    (type)" bullets. Always parses as a table - views/functions/volumes
    never existed in this older transcript layout.
This is the same dual-format logic SQL_2_Fabric.py uses, kept in one place
here so every target generator can share it instead of re-implementing it.
Views/functions/volumes are additive, separate keys rather than being mixed
into "tables" so DB2_2_Fabric.py's existing tables-only Delta table sync
keeps working unchanged.

Medallion layer comes from SECTION 5's "Medallion Architecture Mapping
Table" (Table Name / Layer (Bronze/Silver/Gold) / Reason). Load strategy
comes from SECTION 7's "Incremental Load Strategy" sentence, which names
the Medium/Large tables that should use incremental load; every other
table defaults to Full Load. Both are best-effort: if --migration-plan
isn't supplied (or the file doesn't have these sections), every table
just gets medallion_layer=null / load_strategy=null and DB2_2_Fabric.py
falls back to "(no layer assigned)" / "(not specified)".

Usage
-----
    python plan_to_json.py \\
        [--assessment path/to/Assessment_Report.docx] \\
        [--migration-plan path/to/Migration_Plan.docx] \\
        [--source-system databricks] [--database-name sales_prod] \\
        [--output path/to/migration_plan.json]

Both docx paths default to the generic "Assesment Report.docx" /
"AI_Migration_Plan.docx" copies that Agents_PipeLine.py leaves in
AI_Agent_Pipeline/output/ after every run (it always refreshes those two
regardless of source system, for frontend compatibility). --output
defaults to migration_plan.json in that same folder, which is exactly
where DB2_2_Fabric.py looks for it by default.
"""
import argparse
import json
import re
import sys
from pathlib import Path

import docx

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "AI_Agent_Pipeline" / "output"
DEFAULT_ASSESSMENT_PATH = OUTPUT_DIR / "Assesment Report.docx"
DEFAULT_MIGRATION_PLAN_PATH = OUTPUT_DIR / "AI_Migration_Plan.docx"
DEFAULT_OUTPUT_PATH = OUTPUT_DIR / "migration_plan.json"

BULLET = "•"  # the actual column-list bullet character used by the report template


def parse_assessment_report(doc_path):
    """
    Returns a list of {"schema", "table_name", "columns": [{"name", "data_type"}, ...]}
    dicts, parsed from either report layout (see module docstring).
    """
    doc = docx.Document(str(doc_path))

    tables = []
    current = None
    in_columns_old = False

    for element in doc.element.body:
        if element.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(element, doc)
            text = p.text.strip()
            if not text:
                continue

            # --- OLD FORMAT CHECK ---
            if text.startswith("Table Name:"):
                if current:
                    tables.append(current)
                t_name = text.split(":", 1)[1].strip()
                current = {"schema": "dbo", "table_name": t_name, "type": "table", "columns": [], "details": {}}
                in_columns_old = False
                continue

            if current and text.startswith("Schema:"):
                current["schema"] = text.split(":", 1)[1].strip()
                continue

            if current and text.startswith("- Columns:"):
                in_columns_old = True
                continue

            if current and in_columns_old:
                if text.startswith("-") and not text.startswith(BULLET):
                    in_columns_old = False
                elif text.endswith(":") and not text.startswith(BULLET):
                    in_columns_old = False
                elif text.startswith(BULLET) or text.startswith("-") or text.startswith("*"):
                    col_part = text.lstrip("•-*").strip()
                    if "(" in col_part and ")" in col_part:
                        name_part, type_part = col_part.split("(", 1)
                        col_name = name_part.strip()
                        data_type_raw = type_part.rsplit(")", 1)[0].strip()
                    else:
                        col_name = col_part.strip()
                        data_type_raw = "string"
                    current["columns"].append({"name": col_name, "data_type": data_type_raw})
                continue

            # --- NEW FORMAT CHECK ---
            is_heading2 = bool(p.style and p.style.name.startswith('Heading 2'))
            if (is_heading2 or text.startswith("5.")) and re.match(r"^5\.\d+\s+", text):
                if current:
                    tables.append(current)
                parts = text.split(" ", 1)
                t_name = parts[1].strip()
                current = {"schema": "dbo", "table_name": t_name, "type": "table", "columns": [], "details": {}}
                in_columns_old = False
                continue

            if current:
                # "The <table|view|function|procedure|volume> <schema>.<name>
                # is mapped..." - the intro sentence docx_generator.py writes
                # for every object type under Section 5. Captures both the
                # schema and the object type in one match so plan_to_json can
                # route tables/views/functions/procedures/volumes into
                # separate output lists.
                match = re.search(
                    r"The (table|view|function|procedure|volume) ([a-zA-Z0-9_]+)\." + re.escape(current["table_name"]) + r" is mapped",
                    text,
                    re.IGNORECASE,
                )
                if match:
                    current["type"] = match.group(1).lower()
                    current["schema"] = match.group(2)

        elif element.tag.endswith('tbl') and current:
            t = docx.table.Table(element, doc)
            if len(t.rows) > 0 and len(t.columns) >= 2:
                header_text = t.rows[0].cells[0].text.strip().lower()
                if "column" in header_text:
                    for row in t.rows[1:]:
                        if len(row.cells) >= 2:
                            c_name = row.cells[0].text.strip()
                            c_type = row.cells[1].text.strip()
                            if c_name and c_type:
                                current["columns"].append({"name": c_name, "data_type": c_type})
                elif header_text == "property":
                    # The View/Function/Volume "Object Details" table
                    # (Property/Value) - carries Definition/Return Type/
                    # Volume Type/Storage Location through to migration_plan.json.
                    for row in t.rows[1:]:
                        if len(row.cells) >= 2:
                            prop = row.cells[0].text.strip()
                            val = row.cells[1].text.strip()
                            if prop:
                                current["details"][prop] = val

    if current:
        tables.append(current)

    return tables


def extract_source_system(doc_path):
    """
    Best-effort scan of every 2-column table in the assessment report for a
    "Source System" / <value> row (the "3. Source Database Summary" table
    docx_generator.py always emits). Returns None if not found.
    """
    doc = docx.Document(str(doc_path))
    for element in doc.element.body:
        if not element.tag.endswith('tbl'):
            continue
        t = docx.table.Table(element, doc)
        if len(t.columns) != 2:
            continue
        for row in t.rows:
            if len(row.cells) < 2:
                continue
            key = row.cells[0].text.strip().lower()
            if key == "source system":
                value = row.cells[1].text.strip()
                if value:
                    return value
    return None


def parse_medallion_mapping(doc_path):
    """
    Returns {table_name_lower: layer} parsed from SECTION 5's "Medallion
    Architecture Mapping Table" (header: Table Name / Layer (Bronze/Silver/
    Gold) / Reason). Returns {} if the section/table isn't found.
    """
    doc = docx.Document(str(doc_path))
    mapping = {}
    saw_marker = False

    for element in doc.element.body:
        if element.tag.endswith('p'):
            p = docx.text.paragraph.Paragraph(element, doc)
            text = p.text.strip().lower()
            if "medallion" in text and "mapping" in text:
                saw_marker = True
        elif element.tag.endswith('tbl') and saw_marker:
            t = docx.table.Table(element, doc)
            if len(t.rows) > 0 and len(t.columns) >= 2:
                header = t.rows[0].cells[1].text.strip().lower()
                if "layer" in header:
                    for row in t.rows[1:]:
                        if len(row.cells) >= 2:
                            t_name = row.cells[0].text.strip()
                            layer = row.cells[1].text.strip()
                            if t_name and layer:
                                mapping[t_name.lower()] = layer
                    saw_marker = False  # only take the first matching table

    return mapping


def parse_incremental_tables(doc_path):
    """
    Returns a set of lowercased table names that SECTION 7's "Incremental
    Load Strategy" sentence names as Medium/Large (i.e. should use
    incremental load). Every table not in this set defaults to Full Load.
    Returns an empty set if the sentence isn't found (-> everything Full Load).
    """
    doc = docx.Document(str(doc_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)

    match = re.search(
        r"for medium/large tables\s*\(([^)]*)\)\s*,\s*use incremental load",
        full_text,
        re.IGNORECASE,
    )
    if not match:
        return set()

    names = [n.strip().lower() for n in match.group(1).split(",") if n.strip()]
    return set(names)


def build_plan(assessment_path, migration_plan_path=None, source_system=None, database_name=None):
    assessment_path = Path(assessment_path)
    if not assessment_path.exists():
        raise FileNotFoundError(f"Assessment report not found: {assessment_path}")

    parsed_tables = parse_assessment_report(assessment_path)

    layer_map = {}
    incremental_names = set()
    has_plan = False

    if migration_plan_path is not None:
        migration_plan_path = Path(migration_plan_path)
        if migration_plan_path.exists():
            has_plan = True
            layer_map = parse_medallion_mapping(migration_plan_path)
            incremental_names = parse_incremental_tables(migration_plan_path)
        else:
            print(f"[WARN] Migration plan not found at {migration_plan_path} - "
                  f"medallion_layer/load_strategy will be left unset.", file=sys.stderr)

    if not source_system:
        source_system = extract_source_system(assessment_path)
        if source_system:
            print(f"[INFO] source_system inferred from assessment report: '{source_system}'")

    # Section 5 carries tables, views, functions, procedures, and volumes
    # side by side (see docx_generator.py's flat 5.N numbering across all
    # five types). They're bucketed into separate output lists here -
    # rather than all dumped into "tables" - so DB2_2_Fabric.py's existing
    # tables-only Delta table sync keeps behaving exactly as before;
    # views/functions/procedures/volumes are new, additive keys for
    # whatever consumes them next.
    tables_out = []
    views_out = []
    functions_out = []
    procedures_out = []
    volumes_out = []
    for t in parsed_tables:
        name = t["table_name"]
        schema = t.get("schema") or "dbo"
        obj_type = t.get("type") or "table"
        details = t.get("details") or {}

        if obj_type == "view":
            views_out.append({
                "schema": schema,
                "view_name": name,
                "columns": t["columns"],
                "definition": details.get("Definition"),
            })
        elif obj_type == "function":
            functions_out.append({
                "schema": schema,
                "function_name": name,
                "return_type": details.get("Return Type"),
            })
        elif obj_type == "procedure":
            procedures_out.append({
                "schema": schema,
                "procedure_name": name,
            })
        elif obj_type == "volume":
            volumes_out.append({
                "schema": schema,
                "volume_name": name,
                "volume_type": details.get("Volume Type"),
                "storage_location": details.get("Storage Location"),
            })
        else:
            tables_out.append({
                "schema": schema,
                "table_name": name,
                "medallion_layer": layer_map.get(name.lower()) if has_plan else None,
                "load_strategy": (
                    ("Incremental Load" if name.lower() in incremental_names else "Full Load")
                    if has_plan else None
                ),
                "columns": t["columns"],
            })

    return {
        "meta": {
            "source_system": source_system,
            "database_name": database_name,
        },
        "tables": tables_out,
        "views": views_out,
        "functions": functions_out,
        "procedures": procedures_out,
        "volumes": volumes_out,
    }


def main():
    parser = argparse.ArgumentParser(
        description=__doc__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    parser.add_argument("--assessment", default=None, help="Path to the Assessment Report .docx")
    parser.add_argument("--migration-plan", default=None,
                         help="Path to the AI Migration Plan .docx (optional; supplies medallion_layer/load_strategy)")
    parser.add_argument("--no-migration-plan", action="store_true",
                         help="Skip the migration plan entirely, even if the default one exists")
    parser.add_argument("--source-system", default=None,
                         help="Override source system, e.g. 'databricks', 'sqlserver' (default: read from the report)")
    parser.add_argument("--database-name", default=None,
                         help="Source database name, e.g. 'sales_prod' (not present in the report; supply explicitly if needed)")
    parser.add_argument("--output", default=None, help="Path to write migration_plan.json to")
    args = parser.parse_args()

    assessment_path = Path(args.assessment) if args.assessment else DEFAULT_ASSESSMENT_PATH
    if args.no_migration_plan:
        migration_plan_path = None
    else:
        migration_plan_path = Path(args.migration_plan) if args.migration_plan else DEFAULT_MIGRATION_PLAN_PATH
    output_path = Path(args.output) if args.output else DEFAULT_OUTPUT_PATH

    plan = build_plan(
        assessment_path=assessment_path,
        migration_plan_path=migration_plan_path,
        source_system=args.source_system,
        database_name=args.database_name,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(plan, f, indent=2, ensure_ascii=False)

    n_tables = len(plan["tables"])
    n_cols = sum(len(t["columns"]) for t in plan["tables"])
    n_layered = sum(1 for t in plan["tables"] if t["medallion_layer"])
    n_views = len(plan.get("views", []))
    n_functions = len(plan.get("functions", []))
    n_procedures = len(plan.get("procedures", []))
    n_volumes = len(plan.get("volumes", []))
    print(
        f"[OK] Wrote {output_path} ({n_tables} tables, {n_cols} columns, {n_layered} with a medallion layer assigned, "
        f"{n_views} views, {n_functions} functions, {n_procedures} procedures, {n_volumes} volumes)."
    )


if __name__ == "__main__":
    main()
